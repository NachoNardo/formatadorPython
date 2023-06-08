import os
import docx

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from docx import Document

path_textos = os.getcwd() + "\\textos\\"
path_resultado = os.getcwd() + "\\resultado\\"
dir_arquivos = os.listdir(path_textos)


def inserirParagrafo(text):
    par = novo_doc.add_paragraph()
    run = par.add_run()
    font = run.font
    font.size = Pt(11)
    font.name = 'Arial'
    par.text = text
    par_format = par.paragraph_format
    par_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    par_format.left_indent = Inches(0)
    par_format.space_before = Pt(10)
    par_format.space_after = Pt(0)
    par_format.line_spacing = 1.15

for texto in dir_arquivos:
    texto = texto.replace("/", "")
    path_atual = path_textos + texto
    doc = Document(path_atual)
    novo_doc = Document()
    for paragraph in doc.paragraphs:
        text = paragraph.text
        if(text != ""):
            if text.count(".")<2:
                inserirParagrafo(text)
            else:
                novo_par = text.split(".")
                for par in novo_par:
                    if(par != ""):
                        if(par.startswith(" ")):
                            inserirParagrafo(par[1:] + ".")
                        else:
                            inserirParagrafo(par + ".")
    novo_doc.save(path_resultado + texto)